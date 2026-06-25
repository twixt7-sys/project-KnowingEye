"""Generate UTAUT Word documents from JSON.

Usage:
  pip install python-docx openpyxl
  python "docs/testing(UTAUT)/scripts/generate_documents.py"

Output: docs/testing(UTAUT)/documents/
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font

BASE = Path(__file__).resolve().parents[1]
INST = BASE / "instrument-data"
CH = BASE / "chapter-data"
OUT = BASE / "documents"


def load(rel: str) -> dict:
    with open(BASE / rel, encoding="utf-8") as f:
        return json.load(f)


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def table(doc: Document, headers: list[str], rows: list[list]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = str(h)
    for i, row in enumerate(rows, 1):
        for j in range(len(headers)):
            t.rows[i].cells[j].text = str(row[j] if j < len(row) else "")


def survey_questionnaire_docx() -> None:
    data = load("instrument-data/01_survey_questionnaire.json")
    doc = Document()
    title(doc, "USER ACCEPTANCE TESTING (UAT) SURVEY")
    doc.add_paragraph("Based on the Unified Theory of Acceptance and Use of Technology (UTAUT)")
    doc.add_paragraph(data["system"])
    doc.add_paragraph(data["institution"])
    doc.add_paragraph()
    doc.add_paragraph(data["instructions"])
    doc.add_paragraph("5 – Strongly Agree | 4 – Agree | 3 – Neutral | 2 – Disagree | 1 – Strongly Disagree")
    doc.add_paragraph()

    for part in data["parts"]:
        doc.add_heading(f"PART {part['part']}. {part['title']}", level=1)
        if part.get("intro"):
            doc.add_paragraph(part["intro"])
        if part.get("fields"):
            for f in part["fields"]:
                doc.add_paragraph(f"{f.replace('_', ' ').title()}: _________________________")
        if part.get("items"):
            headers = ["No.", "Statement", "5", "4", "3", "2", "1"]
            rows = [[it["number"], it["statement"], "", "", "", "", ""] for it in part["items"]]
            table(doc, headers, rows)
        if part.get("questions"):
            for i, q in enumerate(part["questions"], 1):
                doc.add_paragraph(f"{i}. {q}")
                doc.add_paragraph("Answer: " + "_" * 60)

    doc.save(OUT / "LCC_OSAS_WBSMS_UTAUT_SurveyQuestionnaire.docx")


def utaut_analysis_report_docx() -> None:
    ch6 = load("chapter-data/07_6_utaut_framework.json")
    ch8 = load("chapter-data/07_8_data_analysis.json")
    doc = Document()
    title(doc, "UTAUT Evaluation Report - OSAS WBSMS")
    doc.add_paragraph("Sections 7.6–7.8 (System Testing and Validation)")
    doc.add_paragraph()

    doc.add_heading(ch6["title"], level=1)
    for p in ch6.get("paragraphs", []):
        doc.add_paragraph(p)
    t = ch6["table_7_6_constructs"]
    doc.add_paragraph(t["caption"])
    table(doc, t["headers"], t["rows"])
    t7 = ch6["table_7_7_respondents"]
    doc.add_paragraph(t7["caption"])
    table(doc, t7["headers"], t7["rows"])

    doc.add_heading(ch8["title"], level=1)
    t8 = ch8["table_7_8_likert"]
    doc.add_paragraph(t8["caption"])
    table(doc, ["Scale", "Value", "WM Range", "Interpretation"], [[r["scale"], r["value"], r["wm_range"], r["interpretation"]] for r in t8["rows"]])

    t9 = ch8["table_7_9_pe"]
    doc.add_paragraph(t9["caption"])
    table(doc, t9["headers"], t9["rows"])

    t10 = ch8["table_7_10_consolidated"]
    doc.add_paragraph(t10["caption"])
    table(doc, t10["headers"], t10["rows"])

    t11 = ch8["table_7_11_pearson"]
    doc.add_paragraph(t11["caption"])
    table(doc, t11["headers"], t11["rows"])

    doc.add_heading("Discussion", level=2)
    for p in ch8.get("discussion_paragraphs", []):
        doc.add_paragraph(p)

    doc.save(OUT / "LCC_OSAS_WBSMS_UTAUT_AnalysisReport.docx")


def raw_responses_template_xlsx() -> None:
    raw = load("results-data/raw_responses.json")
    survey = load("instrument-data/01_survey_questionnaire.json")
    ids = []
    for part in survey["parts"]:
        if "items" in part:
            ids.extend([i["item_id"] for i in part["items"]])

    wb = Workbook()
    ws = wb.active
    ws.title = "RawResponses"
    headers = ["respondent_id", "role", "gender", "age_group", "tech_proficiency"] + ids
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True)
    row = raw.get("template_row", {})
    ws.append([row.get(h, "") for h in headers])
    wb.save(OUT / "LCC_OSAS_WBSMS_UTAUT_RawResponsesTemplate.xlsx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    survey_questionnaire_docx()
    utaut_analysis_report_docx()
    raw_responses_template_xlsx()
    print("Generated documents in:", OUT)


if __name__ == "__main__":
    main()

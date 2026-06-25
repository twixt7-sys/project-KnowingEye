"""Generate Word (.docx) and Excel (.xlsx) from JSON in this directory.

Usage:
  pip install python-docx openpyxl
  python docs/testing/document-data/generate_documents.py

Output: docs/testing/documents/
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parent
OUT = ROOT.parent / "documents"


def load(name: str) -> dict:
    with open(ROOT / name, encoding="utf-8") as f:
        return json.load(f)


def style_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = str(k)
        table.rows[i].cells[1].text = str(v)


def add_matrix_table(doc: Document, headers: list[str], data_rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = str(h)
    for i, row in enumerate(data_rows, start=1):
        for j in range(len(headers)):
            val = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = str(val) if val is not None else ""


def build_test_plan() -> None:
    data = load("01_test_plan.json")
    doc = Document()
    proj = data.get("project", {})
    style_title(doc, "Test Plan - OSAS WBSMS (IEEE 829)")
    doc.add_paragraph(f"Group: {proj.get('group', '')}")
    doc.add_paragraph(f"Standard: {data.get('standard', '')}")
    doc.add_paragraph(proj.get("description", ""))
    doc.add_paragraph()

    s = data["sections"]
    tid = s["test_plan_identifier"]
    doc.add_heading(tid["title"], level=1)
    add_kv_table(
        doc,
        [
            ("Unique ID", tid.get("unique_id", "")),
            ("Level", tid.get("level", "")),
            ("Version", tid.get("version", "")),
            ("Date", tid.get("date", "")),
            ("Authors", tid.get("authors", "")),
        ],
    )
    for rh in tid.get("revision_history", []):
        doc.add_paragraph(f"Rev {rh.get('version')}: {rh.get('date')} - {rh.get('change', '')}")

    intro = s["introduction"]
    doc.add_heading(intro["title"], level=1)
    doc.add_paragraph(intro.get("purpose", ""))
    doc.add_paragraph("References:")
    for ref in intro.get("references", []):
        doc.add_paragraph(ref, style="List Bullet")

    scope = s["scope"]
    doc.add_heading(scope["title"], level=1)
    doc.add_paragraph("In scope:")
    for x in scope.get("in_scope", []):
        doc.add_paragraph(x, style="List Bullet")
    doc.add_paragraph("Out of scope for this cycle:")
    for x in scope.get("out_of_scope_for_this_test_cycle", []):
        doc.add_paragraph(x, style="List Bullet")

    ti = s["test_items"]
    doc.add_heading(ti["title"], level=1)
    for it in ti.get("items", []):
        doc.add_paragraph(
            f"{it.get('name', '')} - {it.get('role', '')} ({it.get('version') or it.get('path', '')})"
        )

    ft = s["features_to_be_tested"]
    doc.add_heading(ft["title"], level=1)
    headers = ["ID", "Feature", "Risk", "Linked cases"]
    rows = [
        [r.get("id"), r.get("feature"), r.get("risk"), r.get("linked_cases")]
        for r in ft.get("rows", [])
    ]
    add_matrix_table(doc, headers, rows)

    fn = s["features_not_to_be_tested"]
    doc.add_heading(fn["title"], level=1)
    for r in fn.get("rows", []):
        doc.add_paragraph(f"{r.get('feature')}: {r.get('reason')}", style="List Bullet")

    ap = s["approach"]
    doc.add_heading(ap["title"], level=1)
    for m in ap.get("methods", []):
        doc.add_paragraph(m, style="List Bullet")
    doc.add_paragraph("Tools: " + ", ".join(ap.get("tools", [])))
    doc.add_paragraph("Metrics: " + ", ".join(ap.get("metrics", [])))

    pf = s["item_pass_fail_criteria"]
    doc.add_heading(pf["title"], level=1)
    for c in pf.get("criteria", []):
        doc.add_paragraph(c, style="List Bullet")

    sr = s["suspension_and_resumption"]
    doc.add_heading(sr["title"], level=1)
    doc.add_paragraph("Suspend if:")
    for x in sr.get("suspend_if", []):
        doc.add_paragraph(x, style="List Bullet")
    doc.add_paragraph("Resume after:")
    for x in sr.get("resume_after", []):
        doc.add_paragraph(x, style="List Bullet")

    td = s["test_deliverables"]
    doc.add_heading(td["title"], level=1)
    for x in td.get("list", []):
        doc.add_paragraph(x, style="List Bullet")

    ts = s["test_tasks_and_schedule"]
    doc.add_heading(ts["title"], level=1)
    add_matrix_table(
        doc,
        ["Task", "Owner", "Due"],
        [[t.get("task"), t.get("owner"), t.get("due")] for t in ts.get("tasks", [])],
    )

    env = s["environmental_needs"]
    doc.add_heading(env["title"], level=1)
    add_kv_table(
        doc,
        [
            ("Hardware/OS", env.get("hardware_os", "")),
            ("Software", env.get("software", "")),
            ("Network", env.get("network", "")),
            ("Test data", env.get("test_data", "")),
        ],
    )

    resp = s["responsibilities"]
    doc.add_heading(resp["title"], level=1)
    add_matrix_table(
        doc,
        ["Role", "Person", "Duty"],
        [[a.get("role"), a.get("person"), a.get("duty")] for a in resp.get("assignments", [])],
    )

    st = s["staffing_and_training"]
    doc.add_heading(st["title"], level=1)
    doc.add_paragraph(st.get("text", ""))

    rc = s["risks_and_contingencies"]
    doc.add_heading(rc["title"], level=1)
    add_matrix_table(
        doc,
        ["Risk", "Mitigation"],
        [[r.get("risk"), r.get("mitigation")] for r in rc.get("rows", [])],
    )

    apv = s["approvals"]
    doc.add_heading(apv["title"], level=1)
    add_matrix_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [
            [r.get("role"), r.get("name"), r.get("signature"), r.get("date")]
            for r in apv.get("sign_off_rows", [])
        ],
    )

    doc.save(OUT / "LCC_OSAS_WBSMS_Doc1_TestPlan.docx")


def build_system_report() -> None:
    d = load("03_system_test_report.json")
    doc = Document()
    style_title(doc, "System Test Report - OSAS WBSMS")
    doc.add_paragraph(f"Group: {d.get('group', '')} | Date: {d.get('report_date', '')}")
    doc.add_paragraph(d.get("executive_summary", ""))

    doc.add_heading("Test environment", level=1)
    te = d.get("test_environment", {})
    add_kv_table(doc, [(k.replace("_", " ").title(), str(v)) for k, v in te.items()])

    fr = d["functional_requirements_traceability"]
    doc.add_heading(fr["title"], level=1)
    rows = fr.get("rows", [])
    add_matrix_table(
        doc,
        ["Requirement", "How verified", "Result", "Evidence"],
        [[x.get("requirement"), x.get("how_verified"), x.get("result"), x.get("evidence")] for x in rows],
    )

    nf = d["non_functional_results"]
    doc.add_heading(nf["title"], level=1)
    for r in nf.get("rows", []):
        doc.add_heading(r.get("category", ""), level=2)
        doc.add_paragraph(f"Objective: {r.get('objective', '')}")
        doc.add_paragraph(f"Method: {r.get('method', '')}")
        doc.add_paragraph(f"Result: {r.get('result', '')}")
        doc.add_paragraph(f"Notes: {r.get('notes', '')}")

    ae = d["automated_evidence"]
    doc.add_heading(ae["title"], level=1)
    doc.add_paragraph(f"Command: {ae.get('command', '')}")
    doc.add_paragraph(f"Last run: {ae.get('last_run_date', '')} | Outcome: {ae.get('outcome', '')}")
    doc.add_paragraph("Tests:")
    for t in ae.get("tests", []):
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph(ae.get("instruction", ""))

    dd = d["defects_during_system_test"]
    doc.add_heading(dd["title"], level=1)
    doc.add_paragraph(dd.get("text", ""))
    doc.add_paragraph(", ".join(dd.get("placeholder_ids", [])))

    oa = d["overall_assessment"]
    doc.add_heading(oa["title"], level=1)
    add_kv_table(
        doc,
        [
            ("Planned manual cases", str(oa.get("planned_manual_cases", ""))),
            ("Executed manual cases", str(oa.get("executed_manual_cases", ""))),
            ("Passed manual cases", str(oa.get("passed_manual_cases", ""))),
            ("Readiness", str(oa.get("readiness_statement", ""))),
        ],
    )

    doc.save(OUT / "LCC_OSAS_WBSMS_Doc5_SystemTestReport.docx")


def build_uat_report() -> None:
    d = load("04_uat_plan_and_report.json")
    doc = Document()
    style_title(doc, "UAT Plan & Report - OSAS WBSMS")
    doc.add_paragraph(f"Group: {d.get('group', '')}")

    cov = d.get("cover", {})
    doc.add_heading("Cover", level=1)
    add_kv_table(doc, [(k.replace("_", " ").title(), str(v)) for k, v in cov.items()])

    s1 = d["section_1_overview_scope"]
    doc.add_heading(s1["title"], level=1)
    doc.add_paragraph(s1.get("overview_paragraph", ""))
    doc.add_paragraph("In scope:")
    for x in s1.get("in_scope_modules", []):
        doc.add_paragraph(x, style="List Bullet")
    doc.add_paragraph(f"Out of scope: {s1.get('out_of_scope', '')}")

    s2 = d["section_2_objectives"]
    doc.add_heading(s2["title"], level=1)
    doc.add_paragraph(f"Primary: {s2.get('primary', '')}")
    doc.add_paragraph("Secondary:")
    for x in s2.get("secondary", []):
        doc.add_paragraph(x, style="List Bullet")

    s3 = d["section_3_methodology_phases"]
    doc.add_heading(s3["title"], level=1)
    doc.add_paragraph(s3.get("methodology", ""))
    add_matrix_table(
        doc,
        ["Phase", "Output"],
        [[c.get("phase"), c.get("output")] for c in s3.get("phases", [])],
    )

    s4 = d["section_4_environment"]
    doc.add_heading(s4["title"], level=1)
    add_matrix_table(
        doc,
        ["Item", "Details"],
        [[r.get("item"), r.get("details")] for r in s4.get("rows", [])],
    )

    s5 = d["section_5_roles"]
    doc.add_heading(s5["title"], level=1)
    add_matrix_table(
        doc,
        ["Role", "Person", "Responsibilities"],
        [[r.get("role"), r.get("person"), r.get("responsibilities")] for r in s5.get("rows", [])],
    )

    s6 = d["section_6_test_cases_summary"]
    doc.add_heading(s6["title"], level=1)
    add_matrix_table(
        doc,
        s6["columns"],
        [
            [
                r.get("uat_id", ""),
                r.get("module", ""),
                r.get("scenario_title", ""),
                r.get("role", ""),
                r.get("priority", ""),
            ]
            for r in s6.get("rows", [])
        ],
    )

    s7 = d["section_7_defect_severity"]
    doc.add_heading(s7["title"], level=1)
    add_matrix_table(
        doc,
        ["Severity", "Definition", "Action"],
        [[r.get("severity"), r.get("definition"), r.get("action")] for r in s7.get("rows", [])],
    )

    s8 = d["section_8_schedule"]
    doc.add_heading(s8["title"], level=1)
    add_matrix_table(
        doc,
        ["Date", "Activity", "Responsible", "Output"],
        [
            [r.get("date"), r.get("activity"), r.get("responsible"), r.get("output")]
            for r in s8.get("rows", [])
        ],
    )

    s9 = d["section_9_entry_exit"]
    doc.add_heading(s9["title"], level=1)
    doc.add_paragraph("Entry criteria:")
    add_matrix_table(
        doc,
        ["Criterion", "Description"],
        [[e.get("criterion"), e.get("description")] for e in s9.get("entry", [])],
    )
    doc.add_paragraph("Exit criteria:")
    add_matrix_table(
        doc,
        ["Criterion", "Description"],
        [[e.get("criterion"), e.get("description")] for e in s9.get("exit", [])],
    )

    s10 = d["section_10_assumptions_risks"]
    doc.add_heading(s10["title"], level=1)
    add_matrix_table(
        doc,
        ["Type", "Detail", "Mitigation"],
        [[r.get("type"), r.get("detail"), r.get("mitigation")] for r in s10.get("rows", [])],
    )

    s11 = d["section_11_execution_results"]
    doc.add_heading(s11["title"], level=1)
    cols = s11["columns"]
    body = [
        [
            r.get("uat_id", ""),
            r.get("scenario", ""),
            r.get("tester_role", ""),
            r.get("date_executed", ""),
            r.get("expected", ""),
            r.get("actual", ""),
            r.get("pass_fail", ""),
            r.get("remarks", ""),
        ]
        for r in s11.get("rows", [])
    ]
    add_matrix_table(doc, cols, body)

    s12 = d["section_12_uat_defects"]
    doc.add_heading(s12["title"], level=1)
    doc.add_paragraph(s12.get("instruction", ""))
    add_matrix_table(doc, s12["columns"], [])

    s13 = d["section_13_feedback"]
    doc.add_heading(s13["title"], level=1)
    add_matrix_table(
        doc,
        s13["columns"],
        [
            [r.get("user_role"), r.get("feedback"), r.get("recommended_action"), r.get("status")]
            for r in s13.get("rows", [])
        ],
    )

    s14 = d["section_14_signoff"]
    doc.add_heading(s14["title"], level=1)
    doc.add_paragraph(s14.get("statement", ""))
    add_matrix_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [
            [s.get("role"), s.get("name_print"), s.get("signature"), s.get("date")]
            for s in s14.get("signatures", [])
        ],
    )

    doc.save(OUT / "LCC_OSAS_WBSMS_Doc6_UATPlanReport.docx")


def build_regression_report() -> None:
    d = load("06_regression_test_report.json")
    doc = Document()
    style_title(doc, "Regression Test Report - OSAS WBSMS")
    doc.add_paragraph(f"Group: {d.get('group', '')} | Date: {d.get('report_date', '')}")

    intro = d["introduction"]
    doc.add_heading(intro["title"], level=1)
    doc.add_paragraph(intro.get("text", ""))

    rs = d["regression_scope"]
    doc.add_heading(rs["title"], level=1)
    doc.add_paragraph(f"Full suite re-run: {rs.get('full_suite_re_run')}")
    doc.add_paragraph("Selected test case IDs:")
    for x in rs.get("selected_test_case_ids", []):
        doc.add_paragraph(x, style="List Bullet")
    doc.add_paragraph("Related defects: " + ", ".join(rs.get("related_defect_ids", [])))

    rt = d["results_table"]
    doc.add_heading(rt["title"], level=1)
    headers = rt["columns"]
    keymap = {
        "Test Case ID": "test_case_id",
        "Previous result": "previous_result",
        "New result": "new_result",
        "Pass/Fail": "pass_fail",
        "Tester": "tester",
        "Date": "date",
        "Notes": "notes",
    }
    body = [[str(r.get(keymap.get(h, h), "")) for h in headers] for r in rt.get("rows", [])]
    add_matrix_table(doc, headers, body)

    ar = d["automated_regression"]
    doc.add_heading(ar["title"], level=1)
    doc.add_paragraph(f"Command: {ar.get('command', '')}")
    doc.add_paragraph(f"Result: {ar.get('last_result', '')}")
    doc.add_paragraph(ar.get("note", ""))

    nd = d["new_defects"]
    doc.add_heading(nd["title"], level=1)
    if not nd.get("rows"):
        doc.add_paragraph("None recorded.")

    sm = d["summary"]
    doc.add_heading(sm["title"], level=1)
    doc.add_paragraph(sm.get("text", ""))

    doc.save(OUT / "LCC_OSAS_WBSMS_Doc8_RegressionTestReport.docx")


def build_summary_report() -> None:
    d = load("07_test_summary_report.json")
    doc = Document()
    style_title(doc, "Test Summary Report - OSAS WBSMS")
    doc.add_paragraph(f"Group: {d.get('group', '')} | Date: {d.get('report_date', '')}")
    doc.add_paragraph(d.get("executive_summary", ""))

    so = d["scope_and_objectives"]
    doc.add_heading(so["title"], level=1)
    doc.add_paragraph(so.get("text", ""))

    m = d["metrics"]
    doc.add_heading(m["title"], level=1)
    add_kv_table(
        doc,
        [
            ("Test cases planned", str(m.get("test_cases_planned", ""))),
            ("Test cases executed", str(m.get("test_cases_executed", ""))),
            ("Overall pass rate %", str(m.get("overall_pass_rate_percent", ""))),
        ],
    )
    doc.add_paragraph("By phase:")
    for row in m.get("by_phase", []):
        doc.add_paragraph(
            f"{row.get('phase')}: executed {row.get('executed')}, passed {row.get('passed')}, rate {row.get('pass_rate')}"
        )
    doc.add_paragraph("Defects:")
    de = m.get("defects", {})
    doc.add_paragraph(f"Total logged: {de.get('total_logged')}")
    for k, v in de.get("severity_breakdown", {}).items():
        doc.add_paragraph(f"  {k}: {v}")
    doc.add_paragraph(f"Fixed: {de.get('fixed')} | Deferred: {de.get('deferred')} | Open: {de.get('remaining_open')}")

    ps = d["phase_summary"]
    doc.add_heading(ps["title"], level=1)
    add_matrix_table(
        doc,
        ["Phase", "Outcome", "Notes"],
        [[r.get("phase"), r.get("outcome"), r.get("notes")] for r in ps.get("rows", [])],
    )

    oi = d["outstanding_issues"]
    doc.add_heading(oi["title"], level=1)
    add_matrix_table(
        doc,
        ["Defect ID", "Justification if deferred", "Risk"],
        [[r.get("defect_id"), r.get("justification_if_deferred"), r.get("risk")] for r in oi.get("rows", [])],
    )

    ll = d["lessons_learned"]
    doc.add_heading(ll["title"], level=1)
    for b in ll.get("bullets", []):
        doc.add_paragraph(b, style="List Bullet")

    rec = d["recommendation"]
    doc.add_heading(rec["title"], level=1)
    doc.add_paragraph(f"Selection: {rec.get('selection', '')}")
    doc.add_paragraph(rec.get("rationale", ""))

    sof = d["sign_off"]
    doc.add_heading(sof["title"], level=1)
    add_matrix_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [[r.get("role"), r.get("name"), r.get("signature"), r.get("date")] for r in sof.get("rows", [])],
    )

    doc.save(OUT / "LCC_OSAS_WBSMS_Doc9_TestSummaryReport.docx")


def xlsx_style_header(ws) -> None:
    header_font = Font(bold=True)
    fill = PatternFill("solid", fgColor="DDDDDD")
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = fill
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        max_len = 10
        for c in col:
            if c.value is not None:
                max_len = max(max_len, min(len(str(c.value)), 55))
        ws.column_dimensions[letter].width = max_len + 2


def build_test_cases_xlsx() -> None:
    d = load("02_test_cases_document.json")
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Execution_Log"
    sa = d["sheet_a_execution_log"]
    headers = sa["columns"]
    key_a = {
        "Test Case ID": "test_case_id",
        "Pass/Fail": "pass_fail",
        "Test Execution Date": "test_execution_date",
        "Responsible Developer": "responsible_developer",
        "Responsible Tester": "responsible_tester",
        "Comment": "comment",
        "Additional Comments (other than QA team)": "additional_comments_non_qa",
    }
    ws1.append(headers)
    for row in sa["rows"]:
        ws1.append([row.get(key_a[h], "") for h in headers])
    xlsx_style_header(ws1)

    ws2 = wb.create_sheet("Detailed_Test_Cases")
    sb = d["sheet_b_detailed_test_cases"]
    h2 = sb["columns"]
    ws2.append(h2)
    for r in sb["rows"]:
        steps = r.get("test_steps", [])
        if isinstance(steps, list):
            steps_s = "\n".join(f"{i + 1}. {s}" for i, s in enumerate(steps))
        else:
            steps_s = str(steps)
        ws2.append(
            [
                r.get("test_case_id", ""),
                r.get("module_feature", ""),
                r.get("test_case_title", ""),
                r.get("preconditions", ""),
                steps_s,
                r.get("test_data", ""),
                r.get("expected_result", ""),
                r.get("actual_result", ""),
                r.get("pass_fail", ""),
                r.get("priority", ""),
                r.get("tester_name", ""),
                r.get("date_executed", ""),
            ]
        )
    xlsx_style_header(ws2)
    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row):
        for c in row:
            c.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(OUT / "LCC_OSAS_WBSMS_Doc2_TestCases.xlsx")


def build_defect_log_xlsx() -> None:
    d = load("05_defect_bug_report_log.json")
    wb = Workbook()
    ws = wb.active
    ws.title = "Defect_Log"
    cols = d["columns"]
    ws.append(cols)
    for def_row in d["defects"]:
        steps = def_row.get("steps_to_reproduce", [])
        if isinstance(steps, list):
            steps_s = "\n".join(f"{i + 1}. {s}" for i, s in enumerate(steps))
        else:
            steps_s = str(steps)
        ws.append(
            [
                def_row.get("defect_id", ""),
                def_row.get("title", ""),
                def_row.get("module_feature", ""),
                def_row.get("severity", ""),
                def_row.get("priority", ""),
                def_row.get("phase_found", ""),
                steps_s,
                def_row.get("expected_result", ""),
                def_row.get("actual_result", ""),
                def_row.get("status", ""),
                def_row.get("assigned_to", ""),
                def_row.get("date_reported", ""),
                def_row.get("date_resolved", ""),
                def_row.get("evidence", ""),
            ]
        )
    xlsx_style_header(ws)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for c in row:
            c.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(OUT / "LCC_OSAS_WBSMS_Doc7_DefectLog.xlsx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_test_plan()
    build_system_report()
    build_uat_report()
    build_regression_report()
    build_summary_report()
    build_test_cases_xlsx()
    build_defect_log_xlsx()
    print(f"Generated documents in: {OUT}")


if __name__ == "__main__":
    main()
